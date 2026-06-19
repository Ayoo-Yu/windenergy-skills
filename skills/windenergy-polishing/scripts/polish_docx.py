#!/usr/bin/env python3
"""Apply Word polishing changes as OOXML tracked changes.

Usage:
    python polish_docx.py INPUT.docx CHANGES.json OUTPUT.docx
    python polish_docx.py INPUT.docx CHANGES.json OUTPUT.docx --color-crossrefs
    python polish_docx.py --list-blocks INPUT.docx
    python polish_docx.py --color-crossrefs INPUT.docx OUTPUT.docx

Changes JSON format:
[
    {
        "paragraph_index": 25,
        "old": "exact text to find in paragraph",
        "new": "replacement text",
        "note": "optional description"
    },
    {
        "block_id": "t0r1c2p0",
        "old": "text in a table cell",
        "new": "replacement text",
        "note": "optional description"
    },
    {
        "paragraph_index": 240,
        "action": "delete_paragraph",
        "note": "author-approved deletion"
    }
]

Supported targets:
  - Body paragraphs by paragraph_index or block_id such as p25.
  - Table-cell paragraphs by block_id such as t0r1c2p0.
  - Automatic unique-text location when neither paragraph_index nor block_id is
    provided and old appears in exactly one supported block.

Unsupported content:
  - Headers, footers, text boxes, comments, footnotes, endnotes, equations,
    fields, citation-manager fields, images, charts, and embedded objects.

Requirements: python-docx, lxml
"""

from __future__ import annotations

import copy
import json
import re
import sys
import zipfile
from collections import defaultdict
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8")

AUTHOR = "Academic Polishing"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
BLUE = "0000FF"
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
CROSSREF_RE = re.compile(
    r"\b(?i:figs?\.?|figures?)\s+\d+[A-Za-z]?(?:\.\d+)?(?:\([A-Za-z0-9]+\))?"
    r"(?:\s*(?:,|and|or)\s*\d+[A-Za-z]?(?:\.\d+)?(?:\([A-Za-z0-9]+\))?)*"
    r"|\b(?i:tables?)\s+(?:\d+[A-Za-z]?|[A-Z]\.\d+)"
    r"(?:\s*(?:,|and|or)\s*(?:\d+[A-Za-z]?|[A-Z]\.\d+))*"
    r"|\b(?i:appendix|appendices)(?:\s+(?:(?i:table)\s+[A-Z]\.\d+|[A-Z]\d*))?"
)
CAPTION_RE = re.compile(
    r"^(?:Fig\.?|Figure|Table)\s+\d+[A-Za-z]?(?:\([A-Za-z0-9]+\))?[\.:]\s+",
    re.IGNORECASE,
)
CITATION_RANGE_DASH = r"(?:-|\u2013|\u2014)"
CITATION_NUMBER = rf"\d+(?:\s*{CITATION_RANGE_DASH}\s*\d+)?"
CITATION_RE = re.compile(rf"\[(?:{CITATION_NUMBER})(?:\s*,\s*{CITATION_NUMBER})*\]")
SAFE_PARAGRAPH_CHILDREN = {qn("w:pPr"), qn("w:r")}
SAFE_RUN_CHILDREN = {qn("w:rPr"), qn("w:t"), qn("w:lastRenderedPageBreak")}
IGNORABLE_BETWEEN_RUN_TAGS = {"proofErr", "bookmarkStart", "bookmarkEnd"}
UNSAFE_LOCAL_TAGS = {
    "drawing",
    "object",
    "pict",
    "OLEObject",
    "imagedata",
    "shape",
    "fldChar",
    "instrText",
    "fldSimple",
    "footnoteReference",
    "endnoteReference",
    "commentReference",
    "hyperlink",
    "smartTag",
    "sdt",
    "delText",
}


@dataclass
class TextBlock:
    block_id: str
    kind: str
    paragraph: Any
    paragraph_index: int | None = None
    table_index: int | None = None
    row_index: int | None = None
    cell_index: int | None = None
    cell_paragraph_index: int | None = None

    @property
    def text(self) -> str:
        return self.paragraph.text or ""


class TrackedPolisher:
    """Apply text replacements to supported Word paragraphs."""

    def __init__(self, doc_path: Path):
        self.doc_path = Path(doc_path)
        self.doc = Document(str(doc_path))
        self._change_id = 100
        self.blocks = self._build_blocks()
        self.blocks_by_id = {block.block_id: block for block in self.blocks}

    def _next_id(self) -> str:
        self._change_id += 1
        return str(self._change_id)

    def _local_name(self, tag: str) -> str:
        return tag.rsplit("}", 1)[-1] if tag.startswith("{") else tag

    def _namespace(self, tag: str) -> str:
        return tag[1:].split("}", 1)[0] if tag.startswith("{") else ""

    def _unsafe_paragraph_reasons(self, paragraph) -> list[str]:
        para_elem = paragraph._element
        reasons: list[str] = []

        for elem in para_elem.iter():
            local = self._local_name(elem.tag)
            namespace = self._namespace(elem.tag)
            if namespace == MATH_NS:
                reasons.append("math object")
            elif local in UNSAFE_LOCAL_TAGS:
                reasons.append(local)

        for child in list(para_elem):
            if child.tag not in SAFE_PARAGRAPH_CHILDREN:
                reasons.append(f"complex paragraph child {self._local_name(child.tag)}")

        for run in para_elem.findall(qn("w:r")):
            for child in list(run):
                if child.tag not in SAFE_RUN_CHILDREN:
                    reasons.append(f"complex run child {self._local_name(child.tag)}")

        return sorted(set(reasons))

    def _build_blocks(self) -> list[TextBlock]:
        blocks: list[TextBlock] = []
        for idx, paragraph in enumerate(self.doc.paragraphs):
            blocks.append(
                TextBlock(
                    block_id=f"p{idx}",
                    kind="body",
                    paragraph=paragraph,
                    paragraph_index=idx,
                )
            )
        for table_index, table in enumerate(self.doc.tables):
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    for cell_para_index, paragraph in enumerate(cell.paragraphs):
                        blocks.append(
                            TextBlock(
                                block_id=(
                                    f"t{table_index}r{row_index}"
                                    f"c{cell_index}p{cell_para_index}"
                                ),
                                kind="table",
                                paragraph=paragraph,
                                table_index=table_index,
                                row_index=row_index,
                                cell_index=cell_index,
                                cell_paragraph_index=cell_para_index,
                            )
                        )
        return blocks

    def list_blocks(self) -> list[dict[str, Any]]:
        rows = []
        for block in self.blocks:
            rows.append(
                {
                    "block_id": block.block_id,
                    "kind": block.kind,
                    "paragraph_index": block.paragraph_index,
                    "table_index": block.table_index,
                    "row_index": block.row_index,
                    "cell_index": block.cell_index,
                    "cell_paragraph_index": block.cell_paragraph_index,
                    "unsupported_content": self._unsafe_paragraph_reasons(block.paragraph),
                    "text_preview": block.text[:160],
                }
            )
        return rows

    def package_risk_summary(self) -> dict[str, Any]:
        """Summarize package parts that automatic paragraph edits must preserve."""
        try:
            with zipfile.ZipFile(self.doc_path) as archive:
                names = archive.namelist()
                xml_names = [name for name in names if name.endswith(".xml")]
                ole_xml_parts: list[str] = []
                math_xml_parts: list[str] = []
                for name in xml_names:
                    text = archive.read(name).decode("utf-8", errors="ignore")
                    if "OLEObject" in text or "oleObject" in text:
                        ole_xml_parts.append(name)
                    if "oMath" in text or MATH_NS in text or "MathType" in text:
                        math_xml_parts.append(name)
        except Exception as exc:
            return {"readable_package": False, "error": str(exc)}

        embeddings = [name for name in names if name.startswith("word/embeddings/")]
        active_x = [name for name in names if name.startswith("word/activeX/")]
        custom_xml = [name for name in names if name.startswith("customXml/")]
        return {
            "readable_package": True,
            "embedding_parts": embeddings,
            "active_x_parts": active_x,
            "custom_xml_parts": custom_xml,
            "ole_xml_parts": ole_xml_parts,
            "math_xml_parts": math_xml_parts,
            "requires_package_preserving_save": bool(
                embeddings or active_x or custom_xml or ole_xml_parts or math_xml_parts
            ),
        }

    def _make_run(self, text: str, rpr=None):
        run = etree.Element(qn("w:r"))
        if rpr is not None:
            run.append(copy.deepcopy(rpr))
        text_elem = etree.SubElement(run, qn("w:t"))
        text_elem.set(XML_SPACE, "preserve")
        text_elem.text = text
        return run

    def _make_del(self, text: str, rpr, date: str):
        deletion = etree.Element(qn("w:del"))
        deletion.set(qn("w:id"), self._next_id())
        deletion.set(qn("w:author"), AUTHOR)
        deletion.set(qn("w:date"), date)
        run = etree.SubElement(deletion, qn("w:r"))
        if rpr is not None:
            run.append(copy.deepcopy(rpr))
        deleted_text = etree.SubElement(run, qn("w:delText"))
        deleted_text.set(XML_SPACE, "preserve")
        deleted_text.text = text
        return deletion

    def _make_ins(self, text: str, rpr, date: str):
        insertion = etree.Element(qn("w:ins"))
        insertion.set(qn("w:id"), self._next_id())
        insertion.set(qn("w:author"), AUTHOR)
        insertion.set(qn("w:date"), date)
        insertion.append(self._make_run(text, rpr))
        return insertion

    def _has_ancestor(self, elem, tag_name: str) -> bool:
        parent = elem.getparent()
        target = qn(tag_name)
        while parent is not None:
            if parent.tag == target:
                return True
            parent = parent.getparent()
        return False

    def _set_run_color(self, run, color: str) -> None:
        rpr = run.find(qn("w:rPr"))
        if rpr is None:
            rpr = etree.Element(qn("w:rPr"))
            run.insert(0, rpr)
        color_elem = rpr.find(qn("w:color"))
        if color_elem is None:
            color_elem = etree.SubElement(rpr, qn("w:color"))
        for attr in ["w:themeColor", "w:themeTint", "w:themeShade"]:
            color_elem.attrib.pop(qn(attr), None)
        color_elem.set(qn("w:val"), color)

    def _clone_run_with_text(self, run, text: str, color: str | None = None):
        clone = copy.deepcopy(run)
        text_elems = clone.findall(qn("w:t"))
        if len(text_elems) != 1:
            return None
        text_elem = text_elems[0]
        text_elem.set(XML_SPACE, "preserve")
        text_elem.text = text
        for marker in clone.findall(qn("w:lastRenderedPageBreak")):
            marker.getparent().remove(marker)
        if color:
            self._set_run_color(clone, color)
        return clone

    def _is_safe_text_run(self, run) -> bool:
        if self._has_ancestor(run, "w:del"):
            return False
        for child in list(run):
            if child.tag not in SAFE_RUN_CHILDREN:
                return False
        text_elems = run.findall(qn("w:t"))
        if len(text_elems) != 1:
            return False
        return True

    def _crossref_matches(self, text: str):
        matches = sorted(
            list(CROSSREF_RE.finditer(text)) + list(CITATION_RE.finditer(text)),
            key=lambda match: match.start(),
        )
        filtered = []
        cursor = -1
        for match in matches:
            if match.start() >= cursor:
                filtered.append(match)
                cursor = match.end()
        return filtered

    def _replace_run_with_segments(self, run, pieces: list[tuple[str, bool]], color: str) -> None:
        parent = run.getparent()
        if parent is None:
            return
        insert_at = parent.index(run)
        for text, colored in pieces:
            if not text:
                continue
            clone = self._clone_run_with_text(run, text, color if colored else None)
            if clone is not None:
                parent.insert(insert_at, clone)
                insert_at += 1
        parent.remove(run)

    def _color_crossrefs_in_run_group(self, runs: list[Any], color: str) -> int:
        full = ""
        spans = []
        for run in runs:
            text_elem = run.find(qn("w:t"))
            text = text_elem.text if text_elem is not None and text_elem.text else ""
            spans.append((run, len(full), len(full) + len(text), text))
            full += text

        matches = self._crossref_matches(full)
        if not matches:
            return 0

        for run, span_start, span_end, text in spans:
            intervals = [
                (max(match.start(), span_start), min(match.end(), span_end))
                for match in matches
                if match.start() < span_end and match.end() > span_start
            ]
            if not intervals:
                continue
            pieces: list[tuple[str, bool]] = []
            cursor = span_start
            for start, end in intervals:
                if start > cursor:
                    pieces.append((text[cursor - span_start : start - span_start], False))
                pieces.append((text[start - span_start : end - span_start], True))
                cursor = end
            if cursor < span_end:
                pieces.append((text[cursor - span_start :], False))
            self._replace_run_with_segments(run, pieces, color)

        return len(matches)

    def _color_crossrefs_in_container(self, container, color: str) -> int:
        colored = 0
        group = []
        for child in list(container):
            if child.tag == qn("w:r") and self._is_safe_text_run(child):
                group.append(child)
                continue
            if self._local_name(child.tag) in IGNORABLE_BETWEEN_RUN_TAGS:
                continue
            if group:
                colored += self._color_crossrefs_in_run_group(group, color)
                group = []
        if group:
            colored += self._color_crossrefs_in_run_group(group, color)
        return colored

    def _text_run_containers(self, paragraph):
        para_elem = paragraph._element
        containers = [para_elem]
        for insertion in para_elem.iter(qn("w:ins")):
            if not self._has_ancestor(insertion, "w:del"):
                containers.append(insertion)
        return containers

    def _main_text_bounds(self) -> tuple[int, int]:
        body_blocks = [block for block in self.blocks if block.kind == "body"]
        start = 0
        end = len(body_blocks)
        for idx, block in enumerate(body_blocks):
            text = block.text.strip().lower()
            if text == "introduction" or text.startswith("1. introduction"):
                start = idx
                break
        for idx, block in enumerate(body_blocks):
            text = block.text.strip().lower()
            if text == "references" or text.startswith("references"):
                end = idx
                break
        return start, end

    def _is_caption_block(self, block: TextBlock) -> bool:
        text = block.text.strip()
        style_name = ""
        try:
            style_name = (block.paragraph.style.name or "").lower()
        except Exception:
            style_name = ""
        return "caption" in style_name or bool(CAPTION_RE.match(text))

    def color_crossrefs(self, color: str = BLUE) -> dict[str, Any]:
        body_blocks = [block for block in self.blocks if block.kind == "body"]
        start, end = self._main_text_bounds()
        colored = 0
        paragraphs = 0
        unsafe_paragraphs_with_safe_runs = 0
        for idx, block in enumerate(body_blocks):
            if idx < start or idx >= end:
                continue
            text = block.text.strip()
            if not text or self._is_caption_block(block):
                continue
            has_unsafe_content = bool(self._unsafe_paragraph_reasons(block.paragraph))
            count = 0
            for container in self._text_run_containers(block.paragraph):
                count += self._color_crossrefs_in_container(container, color)
            if count:
                colored += count
                paragraphs += 1
                if has_unsafe_content:
                    unsafe_paragraphs_with_safe_runs += 1
        return {
            "colored_mentions": colored,
            "paragraphs": paragraphs,
            "unsafe_paragraphs_with_safe_runs": unsafe_paragraphs_with_safe_runs,
            "color": color,
        }

    def _paragraph_runs_and_text(self, paragraph):
        para_elem = paragraph._element
        runs = para_elem.findall(qn("w:r"))
        full = ""
        spans = []
        for run in runs:
            text_elem = run.find(qn("w:t"))
            text = text_elem.text if text_elem is not None and text_elem.text else ""
            spans.append(
                {
                    "elem": run,
                    "text": text,
                    "start": len(full),
                    "end": len(full) + len(text),
                }
            )
            full += text
        return runs, spans, full

    def _apply_multi(self, paragraph, changes_list: list[tuple[str, str]]) -> int:
        if self._unsafe_paragraph_reasons(paragraph):
            return 0

        para_elem = paragraph._element
        runs, spans, full = self._paragraph_runs_and_text(paragraph)
        if not runs:
            return 0

        regions: list[tuple[int, int, str, str]] = []
        for old, new in changes_list:
            search_from = 0
            while True:
                pos = full.find(old, search_from)
                if pos == -1:
                    break
                end = pos + len(old)
                if not any(pos < e and end > s for s, e, _, _ in regions):
                    regions.append((pos, end, old, new))
                    break
                search_from = pos + 1

        if not regions:
            return 0

        regions.sort(key=lambda item: item[0])
        first_span = spans[0] if spans else None
        rpr0 = first_span["elem"].find(qn("w:rPr")) if first_span else None
        base_rpr = copy.deepcopy(rpr0) if rpr0 is not None else None
        ppr = para_elem.find(qn("w:pPr"))
        now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

        new_elems = []
        if ppr is not None:
            new_elems.append(ppr)

        cursor = 0
        for start, end, old, new in regions:
            if start > cursor:
                new_elems.append(self._make_run(full[cursor:start], base_rpr))
            new_elems.append(self._make_del(old, base_rpr, now))
            if new:
                new_elems.append(self._make_ins(new, base_rpr, now))
            cursor = end

        if cursor < len(full):
            new_elems.append(self._make_run(full[cursor:], base_rpr))

        for child in list(para_elem):
            para_elem.remove(child)
        for elem in new_elems:
            para_elem.append(elem)

        return len(regions)

    def delete_paragraph(self, paragraph) -> bool:
        if self._unsafe_paragraph_reasons(paragraph):
            return False

        para_elem = paragraph._element
        runs, _spans, full = self._paragraph_runs_and_text(paragraph)
        if not runs or not full:
            return False

        rpr0 = runs[0].find(qn("w:rPr"))
        base_rpr = copy.deepcopy(rpr0) if rpr0 is not None else None
        ppr = para_elem.find(qn("w:pPr"))
        now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

        new_elems = []
        if ppr is not None:
            new_elems.append(ppr)
        new_elems.append(self._make_del(full, base_rpr, now))

        for child in list(para_elem):
            para_elem.remove(child)
        for elem in new_elems:
            para_elem.append(elem)
        return True

    def _resolve_block(self, change: dict[str, Any]) -> tuple[TextBlock | None, str | None]:
        if change.get("block_id"):
            block = self.blocks_by_id.get(str(change["block_id"]))
            if block is None:
                return None, f"block_id not found: {change['block_id']}"
            return block, None

        if "paragraph_index" in change:
            try:
                idx = int(change["paragraph_index"])
            except (TypeError, ValueError):
                return None, f"invalid paragraph_index: {change.get('paragraph_index')}"
            block = self.blocks_by_id.get(f"p{idx}")
            if block is None:
                return None, f"paragraph_index out of range: {idx}"
            return block, None

        old = change.get("old")
        if old:
            matches = [block for block in self.blocks if old in block.text]
            if len(matches) == 1:
                return matches[0], None
            if not matches:
                return None, "old text not found in supported body or table blocks"
            return None, "old text appears in multiple supported blocks; add block_id"

        return None, "missing block_id, paragraph_index, or old text"

    def apply_all(self, changes: list[dict[str, Any]]):
        applied: list[dict[str, Any]] = []
        not_found: list[dict[str, Any]] = []
        groups: dict[str, list[dict[str, Any]]] = defaultdict(list)
        order: list[str] = []

        for change in changes:
            block, reason = self._resolve_block(change)
            if block is None:
                failed = dict(change)
                failed["reason"] = reason
                not_found.append(failed)
                continue
            if block.block_id not in groups:
                order.append(block.block_id)
            resolved = dict(change)
            resolved["_block"] = block
            groups[block.block_id].append(resolved)

        for block_id in order:
            change_list = groups[block_id]
            block = change_list[0]["_block"]
            paragraph = block.paragraph
            unsafe_reasons = self._unsafe_paragraph_reasons(paragraph)

            delete_changes = [ch for ch in change_list if ch.get("action") == "delete_paragraph"]
            if delete_changes:
                change = delete_changes[0]
                old_preview = block.text[:80]
                if unsafe_reasons:
                    failed = dict(change)
                    failed["reason"] = "unsupported complex content in paragraph: " + ", ".join(unsafe_reasons)
                    not_found.append(failed)
                elif self.delete_paragraph(paragraph):
                    applied.append(
                        {
                            "block_id": block_id,
                            "note": change.get("note", "delete paragraph"),
                            "old_preview": old_preview,
                            "new_preview": "(deleted)",
                        }
                    )
                else:
                    failed = dict(change)
                    failed["reason"] = "paragraph has no supported text to delete"
                    not_found.append(failed)
                for extra in change_list:
                    if extra is not change:
                        failed = dict(extra)
                        failed["reason"] = "skipped because paragraph deletion was requested"
                        not_found.append(failed)
                continue

            text_changes = [
                (ch.get("old", ""), ch.get("new", ""), ch)
                for ch in change_list
                if ch.get("old") and "new" in ch
            ]
            if unsafe_reasons:
                for _old, _new, change in text_changes:
                    failed = dict(change)
                    failed["reason"] = "unsupported complex content in paragraph: " + ", ".join(unsafe_reasons)
                    not_found.append(failed)
                continue

            pairs = [(old, new) for old, new, _ch in text_changes]
            applied_count = self._apply_multi(paragraph, pairs) if pairs else 0

            for idx, (old, new, change) in enumerate(text_changes):
                if idx < applied_count:
                    applied.append(
                        {
                            "block_id": block_id,
                            "note": change.get("note", ""),
                            "old_preview": old[:80],
                            "new_preview": new[:80],
                        }
                    )
                else:
                    failed = dict(change)
                    failed["reason"] = "old text not found in resolved block or overlaps another change"
                    not_found.append(failed)

        return applied, not_found

    def _document_xml_bytes(self) -> bytes:
        return etree.tostring(
            self.doc._element,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )

    def save(self, path: Path) -> None:
        """Save by replacing only word/document.xml in the original package."""
        path = Path(path)
        document_xml = self._document_xml_bytes()
        with zipfile.ZipFile(self.doc_path) as source:
            entries = [(info, source.read(info.filename)) for info in source.infolist()]

        path.parent.mkdir(parents=True, exist_ok=True)
        wrote_document = False
        with zipfile.ZipFile(path, "w") as target:
            for info, data in entries:
                if info.filename == "word/document.xml":
                    data = document_xml
                    wrote_document = True
                target.writestr(info, data)
            if not wrote_document:
                target.writestr("word/document.xml", document_xml)


def validate_tracked_changes(path: Path) -> dict[str, bool]:
    try:
        with zipfile.ZipFile(path) as archive:
            xml_names = [name for name in archive.namelist() if name.startswith("word/") and name.endswith(".xml")]
            xml = "\n".join(archive.read(name).decode("utf-8", errors="ignore") for name in xml_names)
    except Exception:
        return {"readable_docx": False, "has_deletions": False, "has_insertions": False}
    return {
        "readable_docx": True,
        "has_deletions": "<w:del " in xml,
        "has_insertions": "<w:ins " in xml,
    }


def preview(value: Any, limit: int = 80) -> str:
    text = "" if value is None else str(value)
    return text if len(text) <= limit else text[: limit - 3] + "..."


def print_block_inventory(input_path: Path) -> int:
    polisher = TrackedPolisher(input_path)
    print(
        json.dumps(
            {
                "package_risk_summary": polisher.package_risk_summary(),
                "blocks": polisher.list_blocks(),
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


def extract_color(args: list[str]) -> tuple[list[str], str]:
    if "--color" not in args:
        return args, BLUE
    idx = args.index("--color")
    if idx + 1 >= len(args):
        raise ValueError("--color requires a hex value")
    color = args[idx + 1].lstrip("#")
    remaining = args[:idx] + args[idx + 2 :]
    return remaining, color


def print_usage() -> None:
    print("Usage: polish_docx.py INPUT.docx CHANGES.json OUTPUT.docx [--color-crossrefs]")
    print("       polish_docx.py --list-blocks INPUT.docx")
    print("       polish_docx.py --color-crossrefs INPUT.docx OUTPUT.docx [--color 0000FF]")


def main() -> int:
    args = sys.argv[1:]
    try:
        args, crossref_color = extract_color(args)
    except ValueError as exc:
        print(f"Error: {exc}")
        print_usage()
        return 1

    if len(args) == 2 and args[0] == "--list-blocks":
        input_path = Path(args[1])
        if not input_path.exists():
            print(f"Error: Input file not found: {input_path}")
            return 1
        return print_block_inventory(input_path)

    color_crossrefs = False
    if "--color-crossrefs" in args:
        color_crossrefs = True
        args = [arg for arg in args if arg != "--color-crossrefs"]

    if color_crossrefs and len(args) == 2:
        input_path = Path(args[0])
        output_path = Path(args[1])
        if not input_path.exists():
            print(f"Error: Input file not found: {input_path}")
            return 1
        polisher = TrackedPolisher(input_path)
        color_result = polisher.color_crossrefs(crossref_color)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        polisher.save(output_path)
        print(f"Input:    {input_path}")
        print(f"Output:   {output_path}")
        print(f"Colored:  {color_result['colored_mentions']} mention(s) in {color_result['paragraphs']} paragraph(s)")
        if color_result["unsafe_paragraphs_with_safe_runs"]:
            print(f"Complex:  {color_result['unsafe_paragraphs_with_safe_runs']} paragraph(s) had safe text runs colored")
        print(f"Color:    #{color_result['color']}")
        return 0

    if len(args) != 3:
        print_usage()
        return 1

    input_path = Path(args[0])
    changes_path = Path(args[1])
    output_path = Path(args[2])

    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        return 1
    if not changes_path.exists():
        print(f"Error: Changes file not found: {changes_path}")
        return 1

    with changes_path.open("r", encoding="utf-8") as handle:
        changes = json.load(handle)

    print(f"Input:    {input_path}")
    print(f"Changes:  {len(changes)} item(s)")
    print()

    polisher = TrackedPolisher(input_path)
    applied, not_found = polisher.apply_all(changes)
    color_result = {}
    if color_crossrefs:
        color_result = polisher.color_crossrefs(crossref_color)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    polisher.save(output_path)
    validation = validate_tracked_changes(output_path)

    print(f"Output:   {output_path}")
    print(f"Applied:  {len(applied)} / {len(changes)}")
    if color_result:
        print(f"Colored:  {color_result['colored_mentions']} cross-reference or citation mention(s)")
        if color_result["unsafe_paragraphs_with_safe_runs"]:
            print(f"Complex:  {color_result['unsafe_paragraphs_with_safe_runs']} paragraph(s) had safe text runs colored")
    print(f"OOXML:    readable={validation['readable_docx']} deletions={validation['has_deletions']} insertions={validation['has_insertions']}")
    print()

    for idx, item in enumerate(applied, 1):
        print(f"  {idx}. {item['block_id']}: {item['note']}")
        print(f"     - {preview(item['old_preview'])}")
        print(f"     + {preview(item['new_preview'])}")

    if not_found:
        print(f"\nNot applied ({len(not_found)}):")
        for item in not_found:
            target = item.get("block_id", item.get("paragraph_index", "(auto)"))
            old_preview = preview(item.get("old", item.get("action", "")))
            reason = item.get("reason", "not found")
            print(f"  - {target}: {old_preview} [{reason}]")

    print("\nDone. Open in Word to accept or reject changes.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
