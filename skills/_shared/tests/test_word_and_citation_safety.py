from __future__ import annotations

import shutil
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "windenergy-polishing" / "scripts"))
sys.path.insert(0, str(ROOT / "windenergy-citation" / "scripts"))

from polish_docx import TrackedPolisher  # noqa: E402
from windenergy_citation import audit_references  # noqa: E402

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
O_NS = "urn:schemas-microsoft-com:office:office"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


class SkillSafetyTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp = Path(tempfile.mkdtemp(prefix="renewable_skill_test_"))

    def tearDown(self) -> None:
        shutil.rmtree(self.tmp)

    def make_docx_with_fake_ole(self) -> Path:
        docx_path = self.tmp / "math_ole.docx"
        doc = Document()
        doc.add_paragraph("Introduction")
        doc.add_paragraph("Safe old text [1].")
        doc.add_paragraph("Formula paragraph")
        doc.save(docx_path)

        edited = self.tmp / "edited.docx"
        with zipfile.ZipFile(docx_path) as source, zipfile.ZipFile(edited, "w") as target:
            document_xml = etree.fromstring(source.read("word/document.xml"))
            for paragraph in document_xml.findall(".//w:p", namespaces={"w": W_NS}):
                paragraph_text = "".join(
                    node.text or "" for node in paragraph.findall(".//w:t", namespaces={"w": W_NS})
                )
                if paragraph_text == "Formula paragraph":
                    run = etree.SubElement(paragraph, qn("w:r"))
                    obj = etree.SubElement(run, qn("w:object"))
                    ole = etree.SubElement(obj, f"{{{O_NS}}}OLEObject")
                    ole.set(f"{{{R_NS}}}id", "rId999")
                    ole.set("ProgID", "Equation.DSMT4")
            for info in source.infolist():
                data = source.read(info.filename)
                if info.filename == "word/document.xml":
                    data = etree.tostring(
                        document_xml,
                        xml_declaration=True,
                        encoding="UTF-8",
                        standalone=True,
                    )
                target.writestr(info, data)
            target.writestr("word/embeddings/fake-mathtype.bin", b"mathtype-bytes")
        return edited

    def test_polisher_preserves_embedding_and_skips_ole_paragraph(self) -> None:
        input_path = self.make_docx_with_fake_ole()
        output_path = self.tmp / "out.docx"
        polisher = TrackedPolisher(input_path)
        applied, not_applied = polisher.apply_all(
            [
                {
                    "paragraph_index": 1,
                    "old": "Safe old text",
                    "new": "Safe new text",
                    "note": "safe text edit",
                },
                {
                    "paragraph_index": 2,
                    "old": "Formula paragraph",
                    "new": "Changed formula paragraph",
                    "note": "must be skipped",
                },
            ]
        )
        polisher.save(output_path)

        self.assertEqual(len(applied), 1)
        self.assertEqual(len(not_applied), 1)
        self.assertIn("unsupported complex content", not_applied[0]["reason"])

        with zipfile.ZipFile(input_path) as source, zipfile.ZipFile(output_path) as target:
            self.assertEqual(
                source.read("word/embeddings/fake-mathtype.bin"),
                target.read("word/embeddings/fake-mathtype.bin"),
            )
            xml = target.read("word/document.xml").decode("utf-8")
        self.assertIn("<w:del ", xml)
        self.assertIn("<w:ins ", xml)
        self.assertIn("Formula paragraph", xml)
        self.assertNotIn("Changed formula paragraph", xml)

    def test_citation_audit_reports_fail_unchecked_range_and_uncited(self) -> None:
        manuscript = self.tmp / "paper.docx"
        doc = Document()
        doc.add_paragraph("Introduction")
        doc.add_paragraph("Prior work is cited [1-3].")
        doc.add_paragraph("References")
        doc.save(manuscript)

        bib = self.tmp / "refs.bib"
        bib.write_text(
            """
@article{ok,
  title={A correct wind article},
  author={A. Author},
  journal={Renewable Energy},
  year={2020},
  doi={10.1000/ok}
}
@article{wrong,
  title={Wrong local title},
  author={B. Author},
  journal={Renewable Energy},
  year={2021},
  doi={10.1000/wrong}
}
@article{unchecked,
  title={Unfindable title},
  author={C. Author},
  journal={Renewable Energy},
  year={2022}
}
@article{uncited,
  title={Uncited verified title},
  author={D. Author},
  journal={Renewable Energy},
  year={2023},
  doi={10.1000/uncited}
}
""",
            encoding="utf-8",
        )

        def fake_lookup(doi: str) -> dict:
            data = {
                "10.1000/ok": {
                    "title": ["A correct wind article"],
                    "container-title": ["Renewable Energy"],
                    "issued": {"date-parts": [[2020]]},
                    "DOI": "10.1000/ok",
                },
                "10.1000/wrong": {
                    "title": ["Different verified title"],
                    "container-title": ["Renewable Energy"],
                    "issued": {"date-parts": [[2021]]},
                    "DOI": "10.1000/wrong",
                },
                "10.1000/uncited": {
                    "title": ["Uncited verified title"],
                    "container-title": ["Renewable Energy"],
                    "issued": {"date-parts": [[2023]]},
                    "DOI": "10.1000/uncited",
                },
            }
            return data[doi]

        report = audit_references(
            manuscript,
            bib,
            doi_lookup=fake_lookup,
            crossref_search=lambda _query, _rows: [],
            openalex_search=lambda _query, _rows: [],
        )

        self.assertFalse(report["strict_ready"])
        self.assertEqual(report["in_text_citations"], [1, 2, 3])
        self.assertEqual(report["uncited_references"], [4])
        statuses = {item["number"]: item["status"] for item in report["references"]}
        self.assertEqual(statuses[1], "PASS")
        self.assertEqual(statuses[2], "FAIL")
        self.assertEqual(statuses[3], "UNCHECKED")
        self.assertEqual(statuses[4], "FAIL")


if __name__ == "__main__":
    unittest.main()
